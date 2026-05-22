import os
import re
import matplotlib.pyplot as plt
import pandas as pd
import csv
from datetime import datetime
import openpyxl
from docx import Document
from docx.shared import Inches
import docx2md


class Plotter:
    SKIP_KEYWORDS = ["skip", "template", "notes", "archive"]

    def __init__(self, input_dir, output_dir, host=""):
        self.input_dir = input_dir
        self.output_dir = output_dir
        self.host = host.rstrip("/")
        os.makedirs(self.output_dir, exist_ok=True)

    def run(self):
        config_file = os.path.join(self.input_dir, "plot_config.xlsx")
        plot_registry = []

        if not os.path.exists(config_file):
            return plot_registry

        # Use ExcelFile to access all sheet names first
        xl = pd.ExcelFile(config_file)
        doc = Document()
        doc.add_heading('Multi-Structure Interaction Diagram Report', 0)

        fig_count = 1

        for sheet in xl.sheet_names:
            # Skip logic
            if any(key in sheet.lower() for key in self.SKIP_KEYWORDS):
                print(f"Skipping sheet: {sheet}")
                continue

            print(f"\n--- Processing Sheet: {sheet} ---")
            df = pd.read_excel(xl, sheet_name=sheet)

            # Ensure the sheet isn't empty and has required columns
            if df.empty or 'Graph_Title' not in df.columns:
                print(f"Sheet {sheet} is empty or missing 'Graph_Title'. Skipping.")
                continue

            for _, row in df.iterrows():
                #_ es el índice de la fila, iterrows() devuelve (index, row)
                # Basic check to skip empty rows in Excel
                if pd.isna(row['Graph_Title']):
                    continue

                print(f"  Plotting: {row['Graph_Title']}...")
                fname, desc = self.process_and_plot(row, sheet)
                plot_registry.append({'Sheet': sheet, 'File Name': fname, 'Description': desc})

                # Add to Word Doc
                doc.add_picture(os.path.join(self.output_dir, fname), width=Inches(6))

                # Simplified Caption (Safe). Añade un pie de foto de estilo de parrafo "Caption"
                p = doc.add_paragraph()
                try:
                    p.style = 'Caption'
                except Exception:
                    p.style = 'Normal'
                run = p.add_run(f"Figure {fig_count}: {desc}")
                run.bold = True
                p.alignment = 1

                fig_count += 1
                doc.add_page_break()

        # Save CSV
        csv_path = os.path.join(self.output_dir, "Summary_List.csv")
        with open(csv_path, 'w', newline='', encoding='utf-8') as f:
            w = csv.DictWriter(f, fieldnames=['Sheet', 'File Name', 'Description'])
            w.writeheader()
            w.writerows(plot_registry)

        # Guarda archivo Word Doc
        doc_path = os.path.join(self.output_dir, "Final_Report.docx")
        try:
            doc.save(doc_path)
            print(f"\nDone! Results saved in: {self.output_dir}")
        except PermissionError:
            print(f"\n[!] ERROR: Could not save Word file. Please close 'Final_Report.docx' and try again.")

        # Crea copia del archivo Word Doc en archivo MD usando libreria docx2md
        md_str = docx2md.do_convert(doc_path, target_dir=self.output_dir, use_md_table=False)
        md_str = self._absolutize_image_urls(md_str)
        md_path = os.path.join(self.output_dir, "Final_Report.md")
        with open(md_path, "w", encoding="utf-8") as file:
            file.write(md_str)

        return plot_registry

    def _absolutize_image_urls(self, md_str):
        # Rewrite relative `media/imageN.png` paths in the generated markdown
        # to absolute URLs served by the Flask app, e.g.
        # http://localhost:5000/uploads/1/Plots/media/image1.png
        if not self.host:
            return md_str
        parts = self.output_dir.replace(os.sep, "/").split("/")
        if "uploads" not in parts:
            return md_str
        url_path = "/".join(parts[parts.index("uploads"):])  # uploads/1/Plots
        media_base = f"{self.host}/{url_path}/media"
        return re.sub(
            r'src="(?:[^"]*/)?media/([^"]+)"',
            f'src="{media_base}/\\1"',
            md_str,
        )


    def get_excel_data(self, file_path, sheet_name, start_row, end_row, col_x, col_y):
        file_path = str(file_path).strip().replace('"', '').replace("'", "")
        #file_path del plot_config puede ser relativo, se convierte a absoluto
        if not os.path.isabs(file_path):
            file_path = os.path.join(self.input_dir, file_path)
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            ws = wb[sheet_name]
            x, y = [], []
            for row in ws.iter_rows(min_row=start_row, max_row=end_row, values_only=True):
                #values_only=true: convierte en tupla
                if len(row) >= max(col_x, col_y):
                    #asegura que la tupla tiene suficientes valores. Se pasa col_x que es un entero. Si no, daría index error.
                    val_x, val_y = row[col_x - 1], row[col_y - 1]
                    if isinstance(val_x, (int, float)) and isinstance(val_y, (int, float)):
                        x.append(val_x); y.append(val_y)
            return x, y
        except Exception:
            return [], []

    def process_and_plot(self, config_row, sheet_name):
        #config_row: lectura de la fila del archivo plot_config
        #Genera, da formato y guarda un gráfico
        plt.figure(figsize=(10, 7))

        # --- A. ENVELOPE DATA (TWO INDEPENDENT HALVES) ---
        # Fetch Half 1
        dx1, dy1 = self.get_excel_data(
            config_row['Diag_File'], config_row['Diag_Sheet'],
            int(config_row['Diag_1RowStart']), int(config_row['Diag_1RowEnd']),
            int(config_row['Diag_ColM']), int(config_row['Diag_ColP'])
        )

        # Fetch Half 2
        dx2, dy2 = self.get_excel_data(
            config_row['Diag_File'], config_row['Diag_Sheet'],
            int(config_row['Diag_2RowStart']), int(config_row['Diag_2RowEnd']),
            int(config_row['Diag_ColM']), int(config_row['Diag_ColP'])
        )

        # Console Debugging: Check if data is actually different
        if dx1 and dx2:
            print(f"    [Data Check] {config_row['Graph_Title']}:")
            print(f"      Side A X-range: {min(dx1):.1f} to {max(dx1):.1f}")
            print(f"      Side A Y-range: {min(dy1):.1f} to {max(dy1):.1f}")

        # Plotting Envelope (Black lines)
        if dx1 and dy1:
            plt.plot(dx1, dy1, color='k', linewidth=1.5, label=config_row.get('Diag_Label', 'Capacity'), zorder=1)
            #config_row.get('Diag_Label', 'Capacity') es un diccionario. capacity es valor por defecto si campo vacío.

        if dx2 and dy2:
            # Traza la segunda mitad del gráfico, sin laber para no duplicar 'Capacity' en la leyenda
            plt.plot(dx2, dy2, color='k', linewidth=1.5, zorder=1)

        # --- B. DATASET SCATTER PLOTS ---
        colors = ['red', 'green', 'blue', 'orange', 'lime', 'purple']
        for i in range(1, 7):
            p_path = config_row.get(f'File{i}_Path')
            if pd.notna(p_path):
                #true si p_path no es NaN/null
                x, y = self.get_excel_data(
                    p_path, config_row[f'File{i}_Sheet'],
                    int(config_row[f'File{i}_Start']), int(config_row[f'File{i}_End']),
                    int(config_row[f'File{i}_ColX']), int(config_row[f'File{i}_ColY'])
                )
                if x:
                    plt.scatter(x, y, color=colors[i-1], s=12,
                                label=config_row.get(f'File{i}_Label', f'Data {i}'),
                                alpha=0.8, zorder=2)

        # --- C. FORMATTING ---
        plt.xlabel('Bending Moment (kN*m/m)', fontsize=12)
        plt.ylabel('Normal Force (kN/m)', fontsize=12)

        # Title includes sheet name for easy identification
        plt.title(f"{config_row['Graph_Title']}", fontsize=14, fontweight='bold')

        # Flip Y-axis (Standard for Interaction Diagrams). Invierte la dirección del eje Y, en este caso no se usa.
        #plt.gca().invert_yaxis()

        plt.legend(loc='best', fontsize=9, frameon=True)
        plt.grid(True, which='both', linestyle='--', alpha=0.5)

        # --- D. SAVE AND CLEANUP ---
        clean_title = str(config_row['Graph_Title']).strip()
        file_safe_name = "".join([c for c in clean_title if c.isalnum() or c in (' ', '_')]).strip()
        #list comprehension: solo toma carácteres alfanuméricos, espacios o guión bajo.

        # Filename includes sheet to prevent overwriting
        png_name = f"{sheet_name}_{file_safe_name}.png"
        save_path = os.path.join(self.output_dir, png_name)

        plt.tight_layout()
        plt.savefig(save_path, dpi=300)

        # CRITICAL: Fully clear memory/plots before next iteration
        plt.clf()
        plt.close('all')

        description = f"N-M diagram for {file_safe_name}"
        return png_name, description


# --- MAIN EXECUTION ---
if __name__ == "__main__":
    SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
    timestamp = datetime.now().strftime("%Y-%m-%d_%H%M")
    input_dir = os.path.join(SCRIPT_DIR, "inputs")
    output_dir = os.path.join(SCRIPT_DIR, f"Plots_{timestamp}")

    Plotter(input_dir=input_dir, output_dir=output_dir).run()
