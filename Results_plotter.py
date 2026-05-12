import os
import numpy as np
import matplotlib.pyplot as plt
import openpyxl
import pandas as pd
import csv
from datetime import datetime
from docx import Document
from docx.shared import Inches

# --- 1. SETUP GLOBAL PATHS ---
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
timestamp = datetime.now().strftime("%Y-%m-%d_%H%M")
OUTPUT_DIR = os.path.join(SCRIPT_DIR, f"Plots_{timestamp}")

# Keywords to ignore sheets (case-insensitive)
SKIP_KEYWORDS = ["skip", "template", "notes", "archive"]

if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

# --- 2. REUSABLE FUNCTIONS ---

def get_excel_data(file_path, sheet_name, start_row, end_row, col_x, col_y):
    file_path = str(file_path).strip().replace('"', '').replace("'", "")
    if not os.path.isabs(file_path):
        file_path = os.path.join(SCRIPT_DIR, file_path)
    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)
        ws = wb[sheet_name]
        x, y = [], []
        for row in ws.iter_rows(min_row=start_row, max_row=end_row, values_only=True):
            if len(row) >= max(col_x, col_y):
                val_x, val_y = row[col_x - 1], row[col_y - 1]
                if isinstance(val_x, (int, float)) and isinstance(val_y, (int, float)):
                    x.append(val_x); y.append(val_y)
        return x, y
    except Exception:
        return [], []

def process_and_plot(config_row, sheet_name):
    """Generates, formats, and saves a single graph. Symmetry logic removed."""
    plt.figure(figsize=(10, 7))
    
    # --- A. ENVELOPE DATA (TWO INDEPENDENT HALVES) ---
    # Fetch Half 1
    dx1, dy1 = get_excel_data(
        config_row['Diag_File'], config_row['Diag_Sheet'], 
        int(config_row['Diag_Start']), int(config_row['Diag_End']), 
        int(config_row['Diag_ColX']), int(config_row['Diag_ColY'])
    )
    
    # Fetch Half 2
    dx2, dy2 = get_excel_data(
        config_row['Diag_File'], config_row['Diag_Sheet'], 
        int(config_row['Diag_Start']), int(config_row['Diag_End']), 
        int(config_row['Diag_ColX2']), int(config_row['Diag_ColY2'])
    )

    # Console Debugging: Check if data is actually different
    if dx1 and dx2:
        print(f"    [Data Check] {config_row['Graph_Title']}:")
        print(f"      Side A X-range: {min(dx1):.1f} to {max(dx1):.1f}")
        print(f"      Side A Y-range: {min(dy1):.1f} to {max(dy1):.1f}")
     
    # Plotting Envelope (Black lines)
    if dx1 and dy1:
        plt.plot(dx1, dy1, color='k', linewidth=1.5, label=config_row.get('Diag_Label', 'Capacity'), zorder=1)
    
    if dx2 and dy2:
        # We plot the second half without a label to avoid duplicating 'Capacity' in the legend
        plt.plot(dx2, dy2, color='k', linewidth=1.5, zorder=1)

    # --- B. DATASET SCATTER PLOTS ---
    colors = ['red', 'green', 'blue', 'orange','lime','purple']
    for i in range(1, 7):
        p_path = config_row.get(f'File{i}_Path')
        if pd.notna(p_path):
            x, y = get_excel_data(
                p_path, config_row[f'File{i}_Sheet'], 
                int(config_row[f'File{i}_Start']), int(config_row[f'File{i}_End']), 
                int(config_row[f'File{i}_ColX']), int(config_row[f'File{i}_ColY'])
            )
            if x:
                plt.scatter(x, y, color=colors[i-1], s=12, 
                            label=config_row.get(f'File{i}_Label', f'Data {i}'), 
                            alpha=0.8, zorder=2)

    # --- C. FORMATTING ---
    plt.xlabel('Bending Moment (kN*m/m)', fontsize=12)
    plt.ylabel('Normal Force (kN/m)', fontsize=12)
    
    # Title includes sheet name for easy identification
    plt.title(f"{config_row['Graph_Title']}", fontsize=14, fontweight='bold')
    
    # Flip Y-axis (Standard for Interaction Diagrams)
    plt.gca().invert_yaxis()
    
    plt.legend(loc='best', fontsize=9, frameon=True)
    plt.grid(True, which='both', linestyle='--', alpha=0.5)
    
    # --- D. SAVE AND CLEANUP ---
    clean_title = str(config_row['Graph_Title']).strip()
    file_safe_name = "".join([c for c in clean_title if c.isalnum() or c in (' ', '_')]).strip()
    
    # Filename includes sheet to prevent overwriting
    png_name = f"{sheet_name}_{file_safe_name}.png"
    save_path = os.path.join(OUTPUT_DIR, png_name)
    
    plt.tight_layout()
    plt.savefig(save_path, dpi=300)
    
    # CRITICAL: Fully clear memory/plots before next iteration
    plt.clf() 
    plt.close('all')
    
    description = f"N-M diagram for {file_safe_name}"
    return png_name, description
# --- 3. MAIN EXECUTION ---
if __name__ == "__main__":
    config_file = os.path.join(SCRIPT_DIR, "plot_config.xlsx")
    plot_registry = []
    
    if os.path.exists(config_file):
        # Use ExcelFile to access all sheet names first
        xl = pd.ExcelFile(config_file)
        doc = Document()
        doc.add_heading('Multi-Structure Interaction Diagram Report', 0)
        
        fig_count = 1

        for sheet in xl.sheet_names:
            # Skip logic
            if any(key in sheet.lower() for key in SKIP_KEYWORDS):
                print(f"Skipping sheet: {sheet}")
                continue
            
            print(f"\n--- Processing Sheet: {sheet} ---")
            df = pd.read_excel(xl, sheet_name=sheet)
            
            # Ensure the sheet isn't empty and has required columns
            if df.empty or 'Graph_Title' not in df.columns:
                print(f"Sheet {sheet} is empty or missing 'Graph_Title'. Skipping.")
                continue

            for _, row in df.iterrows():
                # Basic check to skip empty rows in Excel
                if pd.isna(row['Graph_Title']): continue

                print(f"  Plotting: {row['Graph_Title']}...")
                fname, desc = process_and_plot(row, sheet)
                plot_registry.append({'Sheet': sheet, 'File Name': fname, 'Description': desc})

                # Add to Word Doc
                doc.add_picture(os.path.join(OUTPUT_DIR, fname), width=Inches(6))
                
                # Simplified Caption (Safe)
                p = doc.add_paragraph()
                try: p.style = 'Caption'
                except: p.style = 'Normal'
                run = p.add_run(f"Figure {fig_count}: {desc}")
                run.bold = True
                p.alignment = 1
                
                fig_count += 1
                doc.add_page_break()

        # Save CSV
        csv_path = os.path.join(OUTPUT_DIR, "Summary_List.csv")
        with open(csv_path, 'w', newline='', encoding='utf-8') as f:
            w = csv.DictWriter(f, fieldnames=['Sheet', 'File Name', 'Description'])
            w.writeheader()
            w.writerows(plot_registry)

        # Save Word Doc
        doc_path = os.path.join(OUTPUT_DIR, "Final_Report.docx")
        try:
            doc.save(doc_path)
            print(f"\nDone! Results saved in: {OUTPUT_DIR}")
            os.startfile(OUTPUT_DIR)
        except PermissionError:
            print(f"\n[!] ERROR: Could not save Word file. Please close 'Final_Report.docx' and try again.")